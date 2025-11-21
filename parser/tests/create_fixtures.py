"""Script to create test fixture DOCX files."""

from pathlib import Path

from docx import Document

fixtures_dir = Path(__file__).parent / "fixtures"
fixtures_dir.mkdir(exist_ok=True)

# Create valid quiz DOCX with labeled choices (A. B. C. D.)
doc = Document()
doc.add_paragraph("Question 1")
doc.add_paragraph("What is the capital of France?")
doc.add_paragraph("A. London")
doc.add_paragraph("B. Paris")
doc.add_paragraph("C. Berlin")
doc.add_paragraph("D. Madrid")
doc.add_paragraph("")  # Empty paragraph (should be filtered)
doc.add_paragraph("Question 2")
doc.add_paragraph("What is 2 + 2?")
doc.add_paragraph("A. 3")
doc.add_paragraph("B. 4")
doc.add_paragraph("C. 5")
doc.add_paragraph("D. 6")
doc.save(fixtures_dir / "valid_quiz.docx")

# Create unlabeled quiz DOCX (like CLD TECH ASSESSMENT format)
doc = Document()
doc.add_paragraph("Question 1")
doc.add_paragraph("What color is the sky?")
doc.add_paragraph("Blue")
doc.add_paragraph("Purple")
doc.add_paragraph("Green")
doc.add_paragraph("Red")
doc.add_paragraph("Question 2")
doc.add_paragraph("What is 2 + 2?")
doc.add_paragraph("3")
doc.add_paragraph("4")
doc.add_paragraph("5")
doc.add_paragraph("6")
doc.save(fixtures_dir / "unlabeled_quiz.docx")

# Create DOCX with empty paragraphs
doc = Document()
doc.add_paragraph("Line 1")
doc.add_paragraph("")
doc.add_paragraph("   ")  # Whitespace only
doc.add_paragraph("Line 2")
doc.save(fixtures_dir / "with_empty_paragraphs.docx")

print(f"Created fixtures in {fixtures_dir}")
